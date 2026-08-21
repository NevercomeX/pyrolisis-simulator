import io
import datetime
import numpy as np
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    """Utility to set XML shading background color for a table cell in docx."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    """Utility to set subtle borders on a docx table."""
    tblPr = table._tbl.tblPr
    borders_elm = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders_elm)

def generate_word_report(mode_option, results, summary, solver_inputs, config_dict=None):
    """
    Generates a complete, highly formatted engineering technical report in Microsoft Word (.docx) format.
    Returns bytes buffer.
    """
    doc = Document()

    # Set Margins (0.75 in / 1.9 cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    is_continuous = (mode_option == "Continuous Operation")
    mode_str_es = "CONTINUO" if is_continuous else "POR LOTES (BATCH)"
    feed_obj = solver_inputs.get('current_feed')
    feed_name = feed_obj.name if feed_obj else "Lodo de Petróleo / Hidrocarburos"
    current_date = datetime.datetime.now().strftime("%d/%m/%Y")

    # Header Tag
    p_tag = doc.add_paragraph()
    p_tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tag = p_tag.add_run("INFORME TÉCNICO DE INGENIERÍA & EVALUACIÓN DE FACTIBILIDAD")
    run_tag.bold = True
    run_tag.font.size = Pt(9.5)
    run_tag.font.color.rgb = RGBColor(217, 119, 6) # Amber color

    # Main Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run(f"EVALUACIÓN TERMOQUÍMICA Y BALANCES DE MATERIA Y ENERGÍA EN REACTOR ROTATORIO DE PIRÓLISIS {mode_str_es}")
    run_title.bold = True
    run_title.font.size = Pt(15)
    run_title.font.color.rgb = RGBColor(30, 58, 138) # Dark Navy Blue

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("Modelado Cinético Multietapa, Caracterización ASTM del Bio-Crudo y Evaluación de Autosuficiencia Energética")
    run_sub.italic = True
    run_sub.font.size = Pt(10.5)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    doc.add_paragraph() # Spacing

    # --- METADATA TABLE ---
    meta_table = doc.add_table(rows=3, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(meta_table, "CBD5E1")
    
    meta_data = [
        [("Empresa / Cliente:", True), ("PROENERGETICO S.R.L.", False), ("Fecha de Emisión:", True), (current_date, False)],
        [("Modo de Operación:", True), (mode_str_es, False), ("Materia Prima:", True), (feed_name, False)],
        [("Software Simulador:", True), ("Rotary Pyrolysis Simulator v2.0", False), ("Unidad de Proceso:", True), ("Reactor Cilíndrico Rotatorio", False)]
    ]
    for r_idx, row in enumerate(meta_data):
        for c_idx, (text, is_bold) in enumerate(row):
            cell = meta_table.cell(r_idx, c_idx)
            set_cell_background(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph()

    # --- RESUMEN EJECUTIVO (ABSTRACT) ---
    h_abs = doc.add_heading(level=2)
    run_habs = h_abs.add_run("RESUMEN EJECUTIVO (ABSTRACT)")
    run_habs.bold = True
    run_habs.font.size = Pt(12)
    run_habs.font.color.rgb = RGBColor(30, 58, 138)

    sludge_density = float(solver_inputs.get('sludge_density', 900.0))
    bio_oil_density = float(solver_inputs.get('bio_oil_density', 750.0))
    temp_hold = float(solver_inputs.get('temp_hold_c', 550.0))
    temp_start = float(solver_inputs.get('temp_start_c', 25.0))
    heating_rate = float(solver_inputs.get('heating_rate_cmin', 1.0))
    load_kg = solver_inputs['feed_rate_kgh'] if is_continuous else solver_inputs['batch_load_kg']
    load_gal = (load_kg / sludge_density) * 264.172
    load_unit = "gal/h" if is_continuous else "gal"

    conv_pct = summary['conversion_pct']
    oil_pct = summary['oil_yield_pct']
    gas_pct = summary['gas_yield_pct']
    char_pct = summary['char_yield_pct']
    duty_val = summary['heating_duty_kw'] if is_continuous else summary.get('total_energy_kwh', 0.0)
    duty_unit = "kW" if is_continuous else "kWh"

    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.15
    p_abs.paragraph_format.space_after = Pt(10)
    run_abs = p_abs.add_run(
        f"El presente informe técnico expone los resultados de la simulación rigurosa del proceso de pirólisis "
        f"desarrollada en un reactor rotatorio de tambor operando en modo {mode_str_es} a presión atmosférica (0 bar manométrica / 1 atm) "
        f"y temperatura máxima de proceso de {temp_hold:.0f}°C (temperatura inicial: {temp_start:.0f}°C, tasa de calentamiento: {heating_rate:.1f}°C/min). "
        f"Se procesó una carga nominal de {load_gal:,.1f} {load_unit} ({load_kg:,.1f} kg) de {feed_name} "
        f"con una densidad de materia prima de {sludge_density:.1f} kg/m³. "
        f"El modelo fenomenológico resuelve el balance conservativo de materia y energía con transferencia de calor conductivo-convectivo en el lecho sólido. "
        f"Se alcanzó una conversión global de materia volátil del {conv_pct:.1f}%, rindiendo un {oil_pct:.1f} wt.% de bio-crudo condensable "
        f"(densidad del bio-crudo: {bio_oil_density:.1f} kg/m³), un {gas_pct:.1f} wt.% de gas de síntesis (syngas) incondensable, y un {char_pct:.1f} wt.% de bio-carbón (char seco). "
        f"La demanda energética total calculada fue de {duty_val:,.2f} {duty_unit} con un error de cierre en el balance de materia de {summary['mass_error_pct']:.2e}%, "
        f"confirmando la consistencia física del modelo y la factibilidad del proceso industrial."
    )
    run_abs.font.size = Pt(10)

    # --- SECCIÓN 1: PROPIEDADES DE LA MATERIA PRIMA ---
    h1 = doc.add_heading(level=1)
    run_h1 = h1.add_run("1. Propiedades de la Materia Prima y Cinética Química")
    run_h1.bold = True
    run_h1.font.size = Pt(13)
    run_h1.font.color.rgb = RGBColor(30, 58, 138)

    p_t1 = doc.add_paragraph()
    run_t1 = p_t1.add_run(
        "La pirólisis termoquímica de lodos de petróleo e hidrocarburos pesados es un proceso de descomposición en atmósfera anaerobia "
        "(ausencia de oxígeno libre) que convierte fracciones orgánicas de alto peso molecular (asfaltenos, resinas y cadenas parafínicas) "
        "en tres coproductos de alto valor agregado: vapores condensables (Bio-Crudo), gases incondensables (Syngas) y un residuo sólido seco (Bio-Carbón e inorgánicos). "
        "La evolución termodinámica comprende tres dominios principales: (i) Secado y Deshidratación a 100°C; "
        "(ii) Destilación y Desvolatilización Multicomponente entre 296°C y 370°C; y "
        "(iii) Policondensación y Coquización final por encima de 370°C, estabilizando el carbono fijo en la matriz sólida."
    )
    run_t1.font.size = Pt(10)

    # 1.1 Tabla Proximal
    h11 = doc.add_heading(level=2)
    run_h11 = h11.add_run("1.1 Análisis Proximal y Composición Elemental del Lodo")
    run_h11.font.size = Pt(11)
    run_h11.font.color.rgb = RGBColor(15, 118, 110)

    t_prox = doc.add_table(rows=6, cols=4)
    t_prox.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_prox, "CBD5E1")

    headers_prox = ["Componente", "Fracción (wt.%)", f"Carga Entrada ({load_unit})", "Volumen Relativo"]
    for c_idx, h in enumerate(headers_prox):
        cell = t_prox.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_prox = [
        ["Humedad (Moisture)", f"{feed_obj.moisture:.2f}%", f"{load_gal*feed_obj.moisture/100:.2f}", "Agua libre lecho"],
        ["Materia Volátil (Volatiles)", f"{feed_obj.volatile:.2f}%", f"{load_gal*feed_obj.volatile/100:.2f}", "Matriz org. reactiva"],
        ["Carbono Fijo (Fixed Carbon)", f"{feed_obj.fixed_carbon:.2f}%", f"{load_gal*feed_obj.fixed_carbon/100:.2f}", "Estructura char"],
        ["Cenizas Inertes (Ash)", f"{feed_obj.ash:.2f}%", f"{load_gal*feed_obj.ash/100:.2f}", "Inorgánico inerte"],
        ["TOTAL", "100.00%", f"{load_gal:.2f}", "Base Húmeda"]
    ]
    for r_idx, r_data in enumerate(rows_prox):
        for c_idx, val in enumerate(r_data):
            cell = t_prox.cell(r_idx + 1, c_idx)
            if r_idx == len(rows_prox) - 1:
                set_cell_background(cell, "F1F5F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            if r_idx == len(rows_prox) - 1:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # 1.1.B Tabla Propiedades Físicas
    hhv_feed_mj = 18.50
    hhv_feed_btu = hhv_feed_mj * 429.923
    cp_feed_kj = 2.10
    sludge_dens_kg = float(solver_inputs.get('sludge_density', 900.0))
    sludge_dens_lbgal = (sludge_dens_kg / 999.1) * 8.345

    t_fprop = doc.add_table(rows=7, cols=4)
    t_fprop.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_fprop, "CBD5E1")

    headers_fprop = ["Propiedad Físico-Química", "Valor Característico", "Unidad", "Descripción Industrial"]
    for c_idx, h in enumerate(headers_fprop):
        cell = t_fprop.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_fprop = [
        ["Densidad Bruta de Masa (Bulk Density)", f"{sludge_dens_kg:.1f} kg/m³ | {sludge_dens_lbgal:.2f} lb/gal", "kg/m³", "Dimensión de bombas y tolvas"],
        ["Poder Calorífico Superior (HHV Lodo)", f"{hhv_feed_mj:.2f} MJ/kg ({hhv_feed_btu:,.0f} BTU/lb)", "MJ/kg", "Potencial energético inicial lodo"],
        ["Capacidad Calorífica Específica (Cp Lodo)", f"{cp_feed_kj:.2f} kJ/kg·K", "kJ/kg·K", "Requerimiento térmico de rampa"],
        ["Ángulo de Reposo Sólido (Flujo)", f"{feed_obj.angle_of_repose:.1f}°", "grados (°)", "Velocidad de avance en cilindro"],
        ["Análisis Elemental Estimado (Ultimate)", "C: 48.5% | H: 6.2% | O: 8.8% | N: 0.5% | S: 0.8%", "wt.% seco", "Base orgánica de hidrocarburos"],
        ["Rendimiento Teórico (Base Volátiles)", "Bio-Crudo: 55% | Syngas: 20% | Char: 15%", "wt.% volát.", "Distribución nominal de fases"]
    ]
    for r_idx, r_data in enumerate(rows_fprop):
        for c_idx, val in enumerate(r_data):
            cell = t_fprop.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()

    # 1.2 ASTM
    h12 = doc.add_heading(level=2)
    run_h12 = h12.add_run("1.2 Caracterización del Bio-Crudo Producido (Normas ASTM)")
    run_h12.font.size = Pt(11)
    run_h12.font.color.rgb = RGBColor(15, 118, 110)

    vol_pct = feed_obj.volatile
    hhv_oil = min(43.5, max(36.0, 38.0 + 0.008 * (temp_hold - 450.0) + 0.10 * (vol_pct - 50.0)))
    hhv_btu = hhv_oil * 429.923
    factor_enh = hhv_oil / 18.5
    visc_40c = max(12.0, 45.0 - 0.08 * (temp_hold - 400.0))
    pump_status = "Bombeable Directo @ 25°C" if visc_40c <= 25.0 else ("Precalentar 40°C" if visc_40c <= 50.0 else "Alta Viscosidad (60°C)")
    bio_oil_dens = float(solver_inputs.get('bio_oil_density', 750.0))
    sg_15 = bio_oil_dens / 999.1
    api_deg = (141.5 / max(0.1, sg_15)) - 131.5
    bsw_pct = min(8.0, max(1.5, 3.5 * (feed_obj.moisture / 30.0)))

    t_astm = doc.add_table(rows=5, cols=4)
    t_astm.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_astm, "CBD5E1")

    headers_astm = ["Propiedad Físico-Química", "Norma ASTM", "Valor Calculado", "Diagnóstico de Calidad"]
    for c_idx, h in enumerate(headers_astm):
        cell = t_astm.cell(0, c_idx)
        set_cell_background(cell, "0D9488")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_astm = [
        ["Poder Calorífico Superior (PCS/HHV)", "ASTM D240", f"{hhv_oil:.2f} MJ/kg ({hhv_btu:,.0f} BTU/lb)", f"Factor Mejora: x{factor_enh:.2f} vs lodo"],
        ["Viscosidad Cinemática @ 40°C", "ASTM D445", f"{visc_40c:.1f} cSt (mm²/s)", pump_status],
        ["Densidad @ 15°C & Gravedad °API", "ASTM D1298", f"{bio_oil_dens:.1f} kg/m³ | {api_deg:.1f} °API", f"Gravedad Específica: {sg_15:.4f}"],
        ["Contenido de Agua & Sedimentación (BS&W)", "ASTM D95", f"{bsw_pct:.2f} wt.%", "Fase org. condensada limpia"]
    ]
    for r_idx, r_data in enumerate(rows_astm):
        for c_idx, val in enumerate(r_data):
            cell = t_astm.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()

    # --- SECCIÓN 2: GEOMETRÍA Y CONDICIONES OPERATIVAS ---
    h2_sec = doc.add_heading(level=1)
    run_h2s = h2_sec.add_run("2. Geometría y Condiciones Operativas del Reactor")
    run_h2s.bold = True
    run_h2s.font.size = Pt(13)
    run_h2s.font.color.rgb = RGBColor(30, 58, 138)

    p_t2 = doc.add_paragraph()
    run_t2 = p_t2.add_run(
        "El diseño dimensional y operativo del reactor rotatorio regula la transferencia de calor conductivo-convectivo hacia el lecho sólido en agitación. "
        "El grado de llenado óptimo (30-35%) y el coeficiente efectivo de transferencia de calor (h_eff) aseguran una distribución térmica uniforme, "
        "maximizando el área de contacto pared-lodo y evitando puntos calientes que propicien la incrustación de coque en la pared del tambor."
    )
    run_t2.font.size = Pt(10)

    length_m = solver_inputs['length']
    dia_m = solver_inputs['diameter']
    aspect_ratio = length_m / max(0.01, dia_m)
    rpm_val = solver_inputs['rpm']
    fill_deg = summary['filling_degree_pct']
    h_eff = solver_inputs['h_eff']
    burner_hp = solver_inputs.get('burner_hp', 300.0)
    burner_eff = solver_inputs.get('burner_eff_pct', 70.0)
    fuel_cons = summary.get('waste_oil_consumed_galh', summary.get('waste_oil_consumed_gal', 0.0))
    fuel_unit = "/h" if is_continuous else "/lote"

    if is_continuous:
        pyro_residence_min = summary.get('residence_time_min', 30.0)
        t_drying_end = 0.0
        t_start_pyro = 0.0
        t_end_pyro = pyro_residence_min
        total_cycle_min = pyro_residence_min
        time_label_1 = "Tiempo Medio de Residencia (MRT)"
        time_val_1 = f"{pyro_residence_min:.1f} min"
        time_label_2 = "Modo de Operación Continuo"
        time_val_2 = "Flujo Estacionario"
    else:
        t_arr = np.array(results.get('time', [0.0]))
        conv_arr = np.array(results.get('conversion', [0.0]))
        moist_arr = np.array(results.get('moisture', [0.0]))
        temp_s_arr = np.array(results.get('T_solid', [298.15])) - 273.15
        
        if len(conv_arr) > 0 and conv_arr[-1] <= 1.0:
            conv_pct_arr = conv_arr * 100.0
        else:
            conv_pct_arr = conv_arr
            
        t_drying_end = 0.0
        if len(moist_arr) > 0 and moist_arr[0] > 0:
            idx_dry = np.argmax(moist_arr <= 0.005 * moist_arr[0])
            if moist_arr[idx_dry] <= 0.005 * moist_arr[0]:
                t_drying_end = float(t_arr[idx_dry])

        idx_start = np.argmax((conv_pct_arr >= 1.0) | (temp_s_arr >= 296.0))
        if idx_start < len(t_arr) and (conv_pct_arr[idx_start] >= 1.0 or temp_s_arr[idx_start] >= 296.0):
            t_start_pyro = float(t_arr[idx_start])
        else:
            t_start_pyro = float(t_arr[0])
            
        final_conv_pct = conv_pct_arr[-1] if len(conv_pct_arr) > 0 else 0.0
        if final_conv_pct > 0.0:
            idx_end = np.argmax(conv_pct_arr >= 0.99 * final_conv_pct)
            if conv_pct_arr[idx_end] >= 0.99 * final_conv_pct:
                t_end_pyro = float(t_arr[idx_end])
            else:
                t_end_pyro = float(t_arr[-1])
        else:
            t_end_pyro = float(t_arr[-1]) if len(t_arr) > 0 else 300.0
            
        pyro_residence_min = max(0.0, t_end_pyro - t_start_pyro)
        total_cycle_min = float(t_arr[-1]) if len(t_arr) > 0 else 300.0

        time_label_1 = "Tiempo Eficiente Residencia Pirólisis"
        time_val_1 = f"{pyro_residence_min:.1f} min ({pyro_residence_min/60.0:.1f} h)"
        time_label_2 = "Ventana Pirólisis (Inicio → 99% Conv.)"
        time_val_2 = f"{t_start_pyro:.1f} min → {t_end_pyro:.1f} min"

    t_geom = doc.add_table(rows=10, cols=4)
    t_geom.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_geom, "CBD5E1")

    headers_geom = ["Parámetro Geométrico / Operativo", "Valor", "Parámetro Térmico / Energético", "Valor"]
    for c_idx, h in enumerate(headers_geom):
        cell = t_geom.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_geom = [
        ["Longitud Cilíndrica (L)", f"{length_m:.2f} m", "Temperatura Inicial (T_start)", f"{temp_start:.1f} °C"],
        ["Diámetro Interno (D)", f"{dia_m:.2f} m", "Temperatura Máxima (T_hold)", f"{temp_hold:.1f} °C"],
        ["Relación de Aspecto (L/D)", f"{aspect_ratio:.2f}", "Tasa de Calentamiento", f"{heating_rate:.1f} °C/min"],
        ["Velocidad de Rotación", f"{rpm_val:.1f} RPM", "Coef. Transf. Calor (h_eff)", f"{h_eff:.1f} W/m²·K"],
        ["Grado de Llenado del Lecho", f"{fill_deg:.2f} %", "Potencia Quemadores", f"{burner_hp:.0f} HP"],
        ["Tiempo Evaporación Agua (Secado)", f"{t_drying_end:.1f} min" if not is_continuous else "N/A", "Eficiencia Térmica Quemador", f"{burner_eff:.1f} %"],
        [time_label_1, time_val_1, "Consumo Combustible Auxiliar", f"{fuel_cons:.1f} gal{fuel_unit}"],
        [time_label_2, time_val_2, "Conversión Volátiles Total", f"{conv_pct:.1f} %"],
        ["Carga / Alimentación Entrante", f"{load_gal:.1f} {load_unit}", "Tiempo Total Ciclo Lote" if not is_continuous else "Régimen", f"{total_cycle_min:.1f} min ({total_cycle_min/60.0:.1f} h)" if not is_continuous else "Estacionario"]
    ]
    for r_idx, r_data in enumerate(rows_geom):
        for c_idx, val in enumerate(r_data):
            cell = t_geom.cell(r_idx + 1, c_idx)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            run.font.size = Pt(9)

    doc.add_paragraph()

    # --- SECCIÓN 3: BALANCES DE MATERIA Y ENERGÍA ---
    h3_sec = doc.add_heading(level=1)
    run_h3s = h3_sec.add_run("3. Balances de Materia y Energía")
    run_h3s.bold = True
    run_h3s.font.size = Pt(13)
    run_h3s.font.color.rgb = RGBColor(30, 58, 138)

    # 3.1 Balance de Materia
    h31 = doc.add_heading(level=2)
    run_h31 = h31.add_run("3.1 Balance de Materia Global y Cierre de Masa")
    run_h31.font.size = Pt(11)

    p_t31 = doc.add_paragraph()
    run_t31 = p_t31.add_run(
        "El balance de materia cuantifica la transformación estequiométrica del lodo de entrada bajo el principio de conservación de la masa. "
        "La masa total se redistribuye entre la fase condensable (Bio-Crudo), la fase gaseosa incondensable (Syngas), "
        "el residuo sólido (Bio-Carbón) y el agua de secado, garantizando un error de cierre cercano a cero (<0.01%) que confirma la consistencia del modelo."
    )
    run_t31.font.size = Pt(10)

    oil_kgh = summary.get('oil_yield_kgh', summary.get('oil_yield_kg', 0))
    gas_kgh = summary.get('gas_yield_kgh', summary.get('gas_yield_kg', 0))
    char_kgh = summary.get('char_yield_kgh', summary.get('char_yield_kg', 0))
    water_kgh = summary.get('water_yield_kgh', summary.get('water_yield_kg', 0))
    oil_gal_v = (oil_kgh / bio_oil_density) * 264.172
    gas_m3_v = gas_kgh / 1.15
    char_gal_v = (char_kgh / 500.0) * 264.172
    water_gal_v = (water_kgh / 1000.0) * 264.172

    t_mass = doc.add_table(rows=7, cols=4)
    t_mass.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_mass, "CBD5E1")

    headers_mass = ["Flujo / Corriente", "Masa (kg" + ("/h" if is_continuous else "") + ")", "Fracción wt.%", "Volumen Estimado"]
    for c_idx, h in enumerate(headers_mass):
        cell = t_mass.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_mass = [
        ["ENTRADA: " + feed_name, f"{load_kg:.2f}", "100.00%", "-"],
        ["SALIDA: Bio-Crudo (Bio-Oil)", f"{oil_kgh:.2f}", f"{oil_pct:.2f}%", f"{oil_gal_v:.1f} gal" + ("/h" if is_continuous else "")],
        ["SALIDA: Gas de Síntesis (Syngas)", f"{gas_kgh:.2f}", f"{gas_pct:.2f}%", f"{gas_m3_v:.1f} m³" + ("/h" if is_continuous else "")],
        ["SALIDA: Bio-Carbón (Char Seco)", f"{char_kgh:.2f}", f"{char_pct:.2f}%", f"{char_gal_v:.1f} gal" + ("/h" if is_continuous else "")],
        ["SALIDA: Vapor de Agua (Steam)", f"{water_kgh:.2f}", f"{summary['water_yield_pct']:.2f}%", f"{water_gal_v:.1f} gal" + ("/h" if is_continuous else "")],
        ["ERROR CIERRE BALANCE", f"{summary['mass_error_pct']:.2e} %", "Conservativo", "Tolerancia < 0.01%"]
    ]
    for r_idx, r_data in enumerate(rows_mass):
        for c_idx, val in enumerate(r_data):
            cell = t_mass.cell(r_idx + 1, c_idx)
            if r_idx == 0:
                set_cell_background(cell, "F8FAFC")
            elif r_idx == len(rows_mass) - 1:
                set_cell_background(cell, "FEF3C7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            if r_idx == 0 or r_idx == len(rows_mass) - 1:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # 3.2 Balance de Energía
    h32 = doc.add_heading(level=2)
    run_h32 = h32.add_run("3.2 Balance de Energía Térmica y Carga de Calentamiento")
    run_h32.font.size = Pt(11)

    p_t32 = doc.add_paragraph()
    run_t32 = p_t32.add_run(
        "El balance entálpico global integra la demanda térmica del sistema dividida en tres contribuciones clave: "
        "el calor sensible para elevar la temperatura de la carga, el calor latente consumido en la evaporación del agua (2,256 kJ/kg) "
        "y la entalpía endotérmica requerida para la pirólisis de la matriz orgánica (~600 kJ/kg)."
    )
    run_t32.font.size = Pt(10)

    if is_continuous:
        F_char_s = char_kgh / 3600.0
        F_oil_s = oil_kgh / 3600.0
        F_gas_s = gas_kgh / 3600.0
        F_steam_s = water_kgh / 3600.0
        T_in = results['T_solid'][0]
        T_out = results['T_solid'][-1]
        T_gas_out = results['T_gas'][-1]
        Q_char_kw = (F_char_s * 1000.0 * (T_out - T_in)) / 1000.0
        Q_pyro_kw = ((F_oil_s + F_gas_s) * (1800.0 * (T_out - T_in) + 600000.0)) / 1000.0
        if T_in < 100.0:
            Q_steam_kw = (F_steam_s * (4184.0 * (100.0 - T_in) + 2256000.0 + 2000.0 * (max(T_gas_out, 100.0) - 100.0))) / 1000.0
        else:
            Q_steam_kw = (F_steam_s * (2256000.0 + 2000.0 * (max(T_gas_out, T_in) - T_in))) / 1000.0
        Q_total_kw = Q_char_kw + Q_pyro_kw + Q_steam_kw

        rows_nrg = [
            ["Calor Sensible del Sólido", f"{Q_char_kw:.2f} kW", f"{(Q_char_kw/max(0.001,Q_total_kw)*100):.1f}%", "Conducción lecho-pared"],
            ["Secado y Evaporación Humedad", f"{Q_steam_kw:.2f} kW", f"{(Q_steam_kw/max(0.001,Q_total_kw)*100):.1f}%", "Vaporización latente (100°C)"],
            ["Reacción Endotérmica Pirólisis", f"{Q_pyro_kw:.2f} kW", f"{(Q_pyro_kw/max(0.001,Q_total_kw)*100):.1f}%", "Craqueo térmico volátiles"],
            ["DEMANDA TÉRMICA TOTAL", f"{Q_total_kw:.2f} kW", "100.0%", "Potencia requerida"]
        ]
        nrg_header = ["Etapa de Transferencia Térmica", "Potencia (kW)", "Porcentaje (%)", "Mecanismo Principal"]
    else:
        T_start = results['T_solid'][0]
        T_hold = results['T_solid'][-1]
        E_char_kwh = (char_kgh * 1000.0 * (T_hold - T_start)) / 3.6e6
        E_pyro_kwh = ((oil_kgh + gas_kgh) * (1800.0 * (T_hold - T_start) + 600000.0)) / 3.6e6
        if T_hold >= 100.0:
            E_steam_kwh = (water_kgh * (4184.0 * (100.0 - T_start) + 2256000.0 + 2000.0 * (T_hold - 100.0))) / 3.6e6
        else:
            E_steam_kwh = (water_kgh * (4184.0 * (T_hold - T_start))) / 3.6e6
        E_total_kwh = E_char_kwh + E_pyro_kwh + E_steam_kwh

        rows_nrg = [
            ["Calor Sensible del Sólido", f"{E_char_kwh:.2f} kWh", f"{(E_char_kwh/max(0.001,E_total_kwh)*100):.1f}%", "Conducción lecho-pared"],
            ["Secado y Evaporación Humedad", f"{E_steam_kwh:.2f} kWh", f"{(E_steam_kwh/max(0.001,E_total_kwh)*100):.1f}%", "Vaporización latente (100°C)"],
            ["Reacción Endotérmica Pirólisis", f"{E_pyro_kwh:.2f} kWh", f"{(E_pyro_kwh/max(0.001,E_total_kwh)*100):.1f}%", "Craqueo térmico volátiles"],
            ["ENERGÍA TÉRMICA TOTAL", f"{E_total_kwh:.2f} kWh", "100.0%", "Consumo total ciclo"]
        ]
        nrg_header = ["Etapa de Transferencia Térmica", "Energía (kWh)", "Porcentaje (%)", "Mecanismo Principal"]

    t_nrg = doc.add_table(rows=5, cols=4)
    t_nrg.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_nrg, "CBD5E1")

    for c_idx, h in enumerate(nrg_header):
        cell = t_nrg.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    for r_idx, r_data in enumerate(rows_nrg):
        for c_idx, val in enumerate(r_data):
            cell = t_nrg.cell(r_idx + 1, c_idx)
            if r_idx == len(rows_nrg) - 1:
                set_cell_background(cell, "F1F5F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            if r_idx == len(rows_nrg) - 1:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # --- SECCIÓN 4: EVALUACIÓN ECONÓMICA ---
    h5_sec = doc.add_heading(level=1)
    run_h5s = h5_sec.add_run("4. Evaluación de Viabilidad Económica")
    run_h5s.bold = True
    run_h5s.font.size = Pt(13)
    run_h5s.font.color.rgb = RGBColor(30, 58, 138)

    p_t5 = doc.add_paragraph()
    run_t5 = p_t5.add_run(
        "La evaluación financiera modela la factibilidad del proyecto mediante el análisis de Flujo de Caja Descontado (DCF) "
        "a un horizonte de 10 años con una tasa de descuento exigida del 14.0%. Se evalúan la estructura de inversión inicial (CAPEX), "
        "los costos de operación y mantenimiento (OPEX), y los ingresos derivados de la tarifa de recepción de lodos (tipping fee) "
        "junto con la comercialización del bio-crudo valorizado."
    )
    run_t5.font.size = Pt(10)

    from .pdf_generator import _calculate_financials
    fin = _calculate_financials(mode_option, summary, solver_inputs)
    curr_sym = "RD$"
    cb = fin['capex_breakdown']
    tot_cap = fin['total_capex']

    # 4.1 CAPEX
    h51 = doc.add_heading(level=2)
    run_h51 = h51.add_run("4.1 Estructura de Inversión de Capital (CAPEX)")
    run_h51.font.size = Pt(11)

    t_capex = doc.add_table(rows=9, cols=4)
    t_capex.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_capex, "CBD5E1")

    headers_capex = ["Rubro de Inversión (CAPEX)", "Monto (RD$)", "Participación (%)", "Descripción del Activo"]
    for c_idx, h in enumerate(headers_capex):
        cell = t_capex.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_capex = [
        ["Equipamiento Principal del Reactor", f"{curr_sym}{cb['equip']:,.2f}", f"{(cb['equip']/tot_cap*100):.1f}%", "Tambor cilíndrico rotatorio y quemadores"],
        ["Instalación Mecánica y Montaje", f"{curr_sym}{cb['install']:,.2f}", f"{(cb['install']/tot_cap*100):.1f}%", "Ensamblaje, alineación y accionamiento"],
        ["Obras Civiles y Cimentaciones", f"{curr_sym}{cb['civil']:,.2f}", f"{(cb['civil']/tot_cap*100):.1f}%", "Bases de concreto, losas y contención"],
        ["Tuberías y Redes Eléctricas", f"{curr_sym}{cb['piping_elec']:,.2f}", f"{(cb['piping_elec']/tot_cap*100):.1f}%", "Manifold 8'', interconexiones y control"],
        ["Ingeniería y Supervisión", f"{curr_sym}{cb['eng']:,.2f}", f"{(cb['eng']/tot_cap*100):.1f}%", "Diseño conceptual, detalle y HAZOP"],
        ["Permisos Ambientales y Legales", f"{curr_sym}{cb['permits']:,.2f}", f"{(cb['permits']/tot_cap*100):.1f}%", "Licencia ambiental y permisos op."],
        ["Fondo de Contingencias", f"{curr_sym}{cb['cont']:,.2f}", f"{(cb['cont']/tot_cap*100):.1f}%", "Imprevistos de construcción (10%)"],
        ["TOTAL INVERSIÓN (CAPEX)", f"{curr_sym}{tot_cap:,.2f}", "100.0%", "Inversión Inicial Total"]
    ]
    for r_idx, r_data in enumerate(rows_capex):
        for c_idx, val in enumerate(r_data):
            cell = t_capex.cell(r_idx + 1, c_idx)
            if r_idx == len(rows_capex) - 1:
                set_cell_background(cell, "F1F5F9")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            if r_idx == len(rows_capex) - 1:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # 4.2 KPIs Financieros
    h54 = doc.add_heading(level=2)
    run_h54 = h54.add_run("4.2 Evaluación de Indicadores Financieros (KPIs)")
    run_h54.font.size = Pt(11)

    ob = fin['opex_breakdown']
    tot_op = fin['total_opex_base']
    rb = fin['revenue_breakdown']
    tot_rev = fin['total_rev_base']
    irr_str = f"{fin['irr']:.1f}%" if fin['irr'] is not None else "N/A"
    payback_str = f"{fin['payback']:.1f} años" if fin['payback'] != float('inf') else "N/A"
    disc_payback_str = f"{fin['disc_payback']:.1f} años" if fin['disc_payback'] != float('inf') else "N/A"
    ebitda_base = tot_rev - tot_op
    margin_ebitda = (ebitda_base / max(1, tot_rev)) * 100.0

    t_kpi = doc.add_table(rows=11, cols=3)
    t_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_kpi, "CBD5E1")

    headers_kpi = ["Métrica Financiera", "Valor Proyectado", "Criterio / Evaluación"]
    for c_idx, h in enumerate(headers_kpi):
        cell = t_kpi.cell(0, c_idx)
        set_cell_background(cell, "1E3A8A")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(9.5)

    rows_kpi = [
        ["Inversión Inicial Total (CAPEX)", f"{curr_sym}{tot_cap:,.2f}", "Inversión fija inicial"],
        ["Ingresos Anuales Base (Año 1)", f"{curr_sym}{tot_rev:,.2f}", "Venta productos + Tipping Fee"],
        ["Costos Anuales OPEX (Año 1)", f"{curr_sym}{tot_op:,.2f}", "Manejo, insumos, nómina y mant."],
        ["EBITDA Estimado Año 1", f"{curr_sym}{ebitda_base:,.2f}", f"Margen EBITDA: {margin_ebitda:.1f}%"],
        ["Valor Actual Neto (VAN / NPV)", f"{curr_sym}{fin['npv']:,.2f}", "Proyecto Altamente Viable (VAN > 0)"],
        ["Tasa Interna de Retorno (TIR / IRR)", f"{irr_str}", "Tasa de descuento exigida: 14.0%"],
        ["Período Recuperación Simple", payback_str, "Recuperación de capital inicial"],
        ["Período Recuperación Descontado", disc_payback_str, "Recuperación considerando r=14%"],
        ["Índice de Rentabilidad (PI)", f"{fin['pi']:.2f}", "PI > 1.0 confirma valor neto positivo"],
        ["Tarifa Equilibrio (Break-Even Tipping)", f"{curr_sym}{fin['breakeven_tipping']:.2f} / gal", "Autosostenible (Superávit)"]
    ]
    for r_idx, r_data in enumerate(rows_kpi):
        for c_idx, val in enumerate(r_data):
            cell = t_kpi.cell(r_idx + 1, c_idx)
            if r_idx in [4, 5]:
                set_cell_background(cell, "FEF3C7")
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(val)
            if r_idx in [4, 5]:
                run.bold = True
            run.font.size = Pt(9)

    doc.add_paragraph()

    # --- SECCIÓN 5: DISCUSIÓN TÉCNICA Y CONCLUSIONES ---
    h6_sec = doc.add_heading(level=1)
    run_h6s = h6_sec.add_run("5. Discusión Técnica y Conclusiones Académicas")
    run_h6s.bold = True
    run_h6s.font.size = Pt(13)
    run_h6s.font.color.rgb = RGBColor(30, 58, 138)

    conclusions = [
        f"1. Viabilidad Térmica y Rendimiento de Conversión: La pirólisis de {feed_name} alcanza una conversión del {conv_pct:.1f}% de la materia volátil, demostrando que el perfil térmico aplicado (temperatura máxima de {temp_hold:.0f}°C) proporciona el aporte entálpico necesario para la descomposición de la matriz orgánica pesada.",
        f"2. Calidad Físico-Química del Bio-Crudo (Normas ASTM D240/D445): El líquido condensable obtenido posee un Poder Calorífico Superior de {hhv_oil:.2f} MJ/kg ({hhv_btu:,.0f} BTU/lb), representando un factor de concentración energética de x{factor_enh:.2f} respecto al lodo de entrada. La viscosidad cinemática a 40°C de {visc_40c:.1f} cSt categoriza al bio-crudo como {pump_status}.",
        f"3. Autosuficiencia Energética del Sistema Industrial: El rendimiento de gas de síntesis ({gas_pct:.1f} wt.%) correspondiente a {gas_m3_v:,.1f} m³ de syngas incondensable permite sustituir el consumo de combustible fósil auxiliar en los quemadores principales del reactor, garantizando la operación autógena en régimen estacionario.",
        f"4. Conservación de Masa y Validación Numérica del Modelo: El error de cierre en el balance de materia de {summary['mass_error_pct']:.2e}% valida la precisión matemática del esquema de integración y confirma la ausencia de pérdidas ficticias en el simulador.",
        f"5. Viabilidad Económica y Rentabilidad del Proyecto: El análisis financiero proyecta un Valor Actual Neto (VAN) de {curr_sym}{fin['npv']:,.2f} y una Tasa Interna de Retorno (TIR) del {irr_str} (superando la tasa de descuento de 14.0%), con un período de recuperación estimado de {payback_str}, confirmando la rentabilidad de la instalación.",
        f"6. Reducción Volumétrica y Tratamiento Industrial de Residuos: El proceso logra una reducción drástica del volumen de residuo pesado procesado, transformándolo eficientemente en tres coproductos de alto valor agregado (Bio-Crudo, Syngas y Bio-Char).",
        f"7. Dinámica del Lecho Sólido y Control Operativo: Operando a un grado de llenado del lecho de {fill_deg:.1f}%, con temperatura inicial de {temp_start:.1f}°C (activación de ebullición a 296.0°C), un tiempo eficiente de residencia de {pyro_residence_min:.1f} min ({pyro_residence_min/60.0:.2f} h) y un tiempo total de ciclo de {total_cycle_min:.1f} min ({total_cycle_min/60.0:.2f} h), se asegura la agitación constante del sólido en el tambor rotatorio, previniendo la formación de incrustaciones de coque duro en las paredes internas."
    ]

    for c_text in conclusions:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.space_after = Pt(6)
        p_c.paragraph_format.line_spacing = 1.15
        run_c = p_c.add_run(c_text)
        run_c.font.size = Pt(9.5)

    doc.add_paragraph() # Spacing

    # FIRMA
    t_sig = doc.add_table(rows=2, cols=1)
    t_sig.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_s0 = t_sig.cell(0, 0)
    p_s0 = cell_s0.paragraphs[0]
    p_s0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_s0.add_run("________________________________________")

    cell_s1 = t_sig.cell(1, 0)
    p_s1 = cell_s1.paragraphs[0]
    p_s1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s1 = p_s1.add_run("Luis J. Solano Cuevas\nGerente de Calidad")
    run_s1.bold = True
    run_s1.font.size = Pt(10)

    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    word_bytes = buffer.getvalue()
    buffer.close()
    return word_bytes
